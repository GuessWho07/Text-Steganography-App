from docx import Document

utf_check = [0x0020, 0x2004, 0x2005, 0x2006, 0x2007, 0x2009, 0x200A, 0x2008, 0x202F]

code_check = chr(0x0020)+chr(0x2004)+chr(0x2005)+chr(0x2006)+chr(0x2007)+chr(0x2009)+chr(0x200A)+chr(0x2008)+chr(0x202F)

code_table = {
    "0000" : chr(0x202F)+chr(0x0020),
    "0001" : chr(0x0020)+chr(0x2004),
    "0010" : chr(0x2004)+chr(0x0020),
    "0011" : chr(0x0020)+chr(0x2005),
    "0100" : chr(0x2005)+chr(0x0020),
    "0101" : chr(0x0020)+chr(0x2006),
    "0110" : chr(0x2006)+chr(0x0020),
    "0111" : chr(0x0020)+chr(0x2007),
    "1000" : chr(0x2007)+chr(0x0020),
    "1001" : chr(0x0020)+chr(0x2009),
    "1010" : chr(0x2009)+chr(0x0020),
    "1011" : chr(0x0020)+chr(0x200A),
    "1100" : chr(0x200A)+chr(0x0020),
    "1101" : chr(0x0020)+chr(0x2008),
    "1110" : chr(0x2008)+chr(0x0020),
    "1111" : chr(0x0020)+chr(0x202F),
}

reversed_code_table = {
    chr(0x202F) + chr(0x0020): "0000",
    chr(0x0020) + chr(0x2004): "0001",
    chr(0x2004) + chr(0x0020): "0010",
    chr(0x0020) + chr(0x2005): "0011",
    chr(0x2005) + chr(0x0020): "0100",
    chr(0x0020) + chr(0x2006): "0101",
    chr(0x2006) + chr(0x0020): "0110",
    chr(0x0020) + chr(0x2007): "0111",
    chr(0x2007) + chr(0x0020): "1000",
    chr(0x0020) + chr(0x2009): "1001",
    chr(0x2009) + chr(0x0020): "1010",
    chr(0x0020) + chr(0x200A): "1011",
    chr(0x200A) + chr(0x0020): "1100",
    chr(0x0020) + chr(0x2008): "1101",
    chr(0x2008) + chr(0x0020): "1110",
    chr(0x0020) + chr(0x202F): "1111"
}

def text_to_binary_string(file_path):
    try:
        with open(file_path, 'r') as file:
            text = file.read()
            binary_data = bytes(text, 'utf-8')
            binary_string = ''.join(format(byte, '08b') for byte in binary_data)
            if len(binary_string) % 4 != 0:
                binary_string += '0' * (len(binary_string) % 4)
            return binary_string
    except FileNotFoundError:
        return "File not found"
    except Exception as e:
        return f"An error occurred: {str(e)}"

def binary_to_spaces(biner):
    spaces = ""
    while biner != "":
        spaces += code_table[biner[:4]]
        biner = biner[4:]
    return spaces

def count_spaces(docx_file_path):
    counter = 0
    doc = Document(docx_file_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for char in run.text:
                if char in code_check:
                    counter += 1
    return counter

def replace_spaces(docx_file_path, biner):
    doc = Document(docx_file_path)
    binary_index = 0

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            new_text = ""
            for char in text:
                if char == ' ':
                    if binary_index < len(biner):
                        new_text += code_table[biner[binary_index:binary_index + 4]]
                        binary_index += 4
                    else:
                        new_text += ' '
                else:
                    new_text += char

            run.text = new_text

    doc.save(docx_file_path)

def read_spaces(docx_file_path):
    decode = ""
    code_check_set = set(code_check)
    doc = Document(docx_file_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for char in run.text:
                if char in code_check_set:
                    decode += char
    

    return decode

def spaces_to_binary(input_string):
    binary_string = ""

    while(input_string != ""):
        if (input_string[:2] in reversed_code_table):
            binary_string += reversed_code_table[input_string[:2]]
        input_string = input_string[2:]
    return binary_string

def binary_to_text(binary_string):
    binary_chunks = [binary_string[i:i+8] for i in range(0, len(binary_string), 8)]
    word_string = "".join([chr(int(chunk, 2)) for chunk in binary_chunks])
    return word_string

def encrypt(txt_file, word_file):
    biner = text_to_binary_string(txt_file)

    if (count_spaces(word_file) < (len(biner)/4)):
        print("Too short word file")
        return

    
    replace_spaces(word_file, biner)
    return


def decrypt(word_file):
    decode = read_spaces(word_file)
    bindec = spaces_to_binary(decode)
    finito = binary_to_text(bindec)
    print(finito)
    return