from docx import Document
from docx.shared import RGBColor
from math import floor
from docx.shared import Pt
from lxml import etree
import re
from docx.text.run import Run
from docx.oxml.text.run import CT_R

def covert_string_to_bits(text):
    bits = ""
    for char in text:
        ascii_code = ord(char)
        binary_code = bin(ascii_code)[2:].zfill(8)
        bits += binary_code
    return bits

def convert_bits_to_string(bits):
    result = ""
    for i in range(0, len(bits), 8):
        byte = bits[i:i+8]
        decimal_value = int(byte, 2)
        result += chr(decimal_value)
    return result

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    return r, g, b

def calculate_doc_potential(document,depth):
    plaintext=""
    for para in document.paragraphs:
        plaintext += para.text + "\n"
    space_count=0

    for para in document.paragraphs:
        space_count += para.text.count(' ')

    potential = floor((space_count*depth*3)/8)
    print("In this document, you can hide {0} characters".format(potential))

    return potential

def get_modified_color(original_color, stegano_value, number_of_bytes):
    binary_oc = bin(original_color)[2:10-number_of_bytes]
    binary_sv = bin(stegano_value)[2:]
    binary_sv = bin(0)[2:]*(number_of_bytes-len(binary_sv))+binary_sv
    return binary_oc+binary_sv

def get_modified_color_from_bits(original_color, stegano_value_bits, number_of_bytes):
    binary_oc = bin(original_color)[2:10-number_of_bytes]
    binary_sv = stegano_value_bits
    binary_sv = bin(0)[2:]*(number_of_bytes-len(binary_sv))+binary_sv
    return binary_oc+binary_sv

def print_run_style(run):
    print("Run boldness:",run.bold)
    print("Run italic:",run.italic)
    print("Run underline:",run.underline)
    print("Run font:",run.font.name)
    print("Run fontsize:",run.font.size)
    print("Run color:",run.font.color.rgb)

def copy_run_style(source_run, target_run):
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.color.rgb = source_run.font.color.rgb

def get_run_style(run,force=False):
    style = {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "font": run.font.name,
        "fontsize": run.font.size,
        "color": run.font.color.rgb
    }

    if force == True:
        if style["bold"] == None:
            style["bold"] = False

        if style["italic"] == None:
            style["italic"] = False

        if style["underline"] == None:
            style["underline"] = False

        if style["font"] == None:
            style["font"] = "Calibri"

        if style["fontsize"] == None:
            style["fontsize"] = Pt(12)

        if style["color"] == None:
            style["color"] = RGBColor(255,255,255)
        else:
            style["color"] = RGBColor(hex_to_rgb(style["color"]))
    return style

def set_run_style(run, style):
    run.bold = style["bold"]
    run.italic = style["italic"]
    run.underline = style["underline"]
    run.font.name = style["font"]
    run.font.size = style["fontsize"]
    run.font.color.rgb = style["color"]

def check_colors_validity(color_list, paragraph):

    colors_checked = 0
    errors = 0
    for run in paragraph.runs:
        if run.text is ' ' and colors_checked < len(color_list):
            r,g,b = hex_to_rgb(str(run.font.color.rgb))
            valid_r = color_list[colors_checked][0]
            valid_g = color_list[colors_checked][1]
            valid_b = color_list[colors_checked][2]

            if r == valid_r and g == valid_g and b == valid_b:
                errors += 0
            else:
                errors += 1

            colors_checked+=1

    return errors

def change_space_color(doc_path, new_doc_path,colors,to_new=False):
    doc = Document(doc_path)
    placeholder = Document()
    if to_new == True:
        doc_new = Document()

    for triplet in colors:
        while True:
            if len(triplet)<3:
                triplet.append(0)
            elif len(triplet) == 3:
                break

    color_index = 0
    spaces_needed = len(colors)
    spaces_divided = 0

    for para_index, paragraph in enumerate(doc.paragraphs):
        runs_in_para = len(paragraph.runs)
        if to_new == True:
            doc_new.add_paragraph()

        for run_index,run in enumerate(paragraph.runs):
            runs_list = []
            if to_new == True:
                run_style = get_run_style(run,True)
            else:
                run_style = get_run_style(run)

            space_count = 0
            space_pos = []

            for i in range(len(run.text)):
                if run.text[i] == ' ':
                    space_count += 1
                    space_pos.append(i)

            run_text = run.text
            pointer = 0
            split_run_text = re.split(r'(\s+)',run_text)

            split_run_text = [element for element in split_run_text if element != '']
            for text in split_run_text:
                runs_list.append(placeholder.add_paragraph().add_run(text))

            if to_new == False:
                for run_part_index, run_part in enumerate(runs_list):
                    paragraph.add_run(run_part.text)
                    copy_run_style(run, paragraph.runs[-1])
            else:
                for run_part_index, run_part in enumerate(runs_list):
                    doc_new.paragraphs[para_index].add_run(run_part.text)
                    set_run_style(doc_new.paragraphs[para_index].runs[run_part_index], run_style)

            if to_new == False:
                for index,run in enumerate(paragraph.runs):
                    if run.text == ' ' and color_index < len(colors):
                        r, g, b = colors[color_index]
                        run.font.color.rgb = RGBColor(r,g,b)
                        color_index += 1
            else:
                for index,run in enumerate(doc_new.paragraphs[para_index].runs):
                    if run.text == ' ' and color_index < len(colors):
                        r, g, b = colors[color_index]
                        run.font.color.rgb = RGBColor(r,g,b)
                        color_index += 1

        color_validity = check_colors_validity(colors,paragraph)
        if  color_validity == 0:
            print("No colors are invalid")
        else:
            print("Some colors are invalid")

        while runs_in_para != 0:
            p = paragraph._p
            p.remove(paragraph.runs[runs_in_para-1]._r)
            runs_in_para -= 1

    if to_new == True:
        doc_new.save(new_doc_path)
    else:
        doc.save(new_doc_path)

    doc_read = Document(new_doc_path)

def hide_message(message,document,depth):
    message_bytes = covert_string_to_bits(message)
    message_bytes = [message_bytes[i:i+depth*3] for i in range(0, len(message_bytes),depth*3)]

    if len(message_bytes[-1]) == depth*3:
        if message_bytes[-1][-1] == '0':
            message_bytes.append('1' * (depth*3))
        else:
            message_bytes.append('0' * (depth*3))
    else:
        missing_bytes = (depth*3) - len(message_bytes[-1])
        if message_bytes[-1][-1] == '0':
            message_bytes[-1] = message_bytes[-1] + '1' * missing_bytes
        else:
            message_bytes[-1] = message_bytes[-1] + '0' * missing_bytes

    message_bytes = ''.join(message_bytes)
    binary_depth = bin(depth)[2:]

    divided_message = []
    for i in range(0, len(message_bytes),3*depth):
        divided_message.append(message_bytes[i:i+(3*depth)])

    spaces_rgbs = []
    spaces_rgbs.append(get_modified_color(255,depth,4))

    for supergroup in divided_message:
        R = supergroup[:depth]
        G = supergroup[depth:2*depth]
        B = supergroup[2*depth:]
        spaces_rgbs.append(get_modified_color_from_bits(255,R,depth))
        spaces_rgbs.append(get_modified_color_from_bits(255,G,depth))
        spaces_rgbs.append(get_modified_color_from_bits(255,B,depth))

    spaces_rgbs_dec = [int(binary, 2) for binary in spaces_rgbs]
    spaces_rgbs_dec_triplet = [spaces_rgbs_dec[i:i+3] for i in range(0, len(spaces_rgbs_dec), 3)]
    change_space_color(document,"space_color_examples/out.docx",spaces_rgbs_dec_triplet,False)

    return spaces_rgbs

def show_message(hidden_message_list):
    ## ODWRÓC PROCES - DLA TESTÓW
    dec_depth = int(hidden_message_list[0][4:],2)
    hidden_message_list.pop(0)
    decoding_string = ""
    for element in hidden_message_list:
        decoding = element[8-dec_depth:]
        decoding_string += decoding

    decoding_list = [decoding_string[i:i+8] for i in range(0, len(decoding_string), 8)]
    decrypted_message = ""
    for letter in decoding_list:
        decrypted_message+=chr(int(letter,2))

    return decrypted_message

def truncate_binary_string(binary_str):
    last_bit = binary_str[-1]
    target_bit = '0' if last_bit == '0' else '1'

    while True:
        if binary_str[-1] == target_bit:
            binary_str = binary_str[0:len(binary_str)-1]
        else:
            break
    
    return binary_str

def show_message_from_file(document_path):
    colors = []
    run_texts = []
    doc = Document(document_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run_texts.append(run.text)
            if run.text == ' ':
                if run.font.color.rgb != None:
                    colors.append(run.font.color.rgb)

    translated_colors = [hex_to_rgb(str(color)) for color in colors]
    binary_first_color = bin(colors[0][0])[6:]
    depth = int(binary_first_color,2)

    trans_color_stream = []
    for i, triplet in enumerate(translated_colors):
        if i == 0:
            trans_color_stream.append(triplet[1])
            trans_color_stream.append(triplet[2])
        else:
            trans_color_stream.append(triplet[0])
            trans_color_stream.append(triplet[1])
            trans_color_stream.append(triplet[2])

    binary_color_stream = []
    for color in trans_color_stream:
        if bin(color) != '0b0':
            binary_color_stream.append(bin(color)[10-depth:])
        else:
            binary_zero = '0b00000000'
            binary_color_stream.append(binary_zero[10-depth:])

    bin_stream = ""
    for item in binary_color_stream:
        bin_stream += item

    bin_stream = truncate_binary_string(bin_stream)
    characters_bin = [bin_stream[i:i+8] for i in range(0, len(bin_stream), 8)]
    decrypted_message = ""
    for letter in characters_bin:
        decrypted_message+=chr(int(letter,2))

    return decrypted_message


# Main is only executed if the script is called from the command line
def main():
    while True:
        print("Please, choose action")
        action = int(input("1 - Hide message in document\n2 - Extract message from document\n Your choice...: "))
        if action == 1:
            full_path = "C:\\Users\paluc\OneDrive\Dokumenty\\a.docx"

            while True:
                while True:
                    depth = int(input("Please provide depth of the message (range from 1 to 8). Lower values will make the font color less distinghushible from color black, but also less characters will be coded within one space\nYour choice (1-8): "))
                    if depth in range (1,9):
                        break
                    else:
                        print("Invalid value, must be in range (1,8)")
                        continue
                
                # Calculate max length of message that can be hidden in this document
                document = Document(full_path)
                potential = calculate_doc_potential(document,depth)
                message = input("Please, provide message you want to hide in your document\nMessage: ")
                if len(message)>potential:
                    print("Your message is too long for your doc, please pick other parameters...")
                    continue
                else:
                    break
            
            print("For presentation purposes, now the message will be unveiled...")
            hidden_message = hide_message(message,full_path,depth)
            encrypted_message = show_message(hidden_message)

            print("Original message:",message)
            print("Decoded message:",encrypted_message)

        elif action == 2:
            full_path = "C:\\Users\paluc\OneDrive\Dokumenty\\out.docx"
            show_message_from_file(full_path)
        else:
            print("Wrong action has been chosen...")
            continue

if __name__ == "__main__":
    main()
